import openpyxl
from docx import Document
import os

downloads = r"C:\Users\diego\Downloads"

# 1. FICHA ULTRA COMPLETA
print("=" * 80)
print("DOCUMENTO 1: FICHA ULTRA COMPLETA.xlsx")
print("=" * 80)
wb = openpyxl.load_workbook(os.path.join(downloads, "FICHA ULTRA COMPLETA..xlsx"))
for sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
    print(f"\n--- Hoja: {sheet_name} ---")
    for row in ws.iter_rows(values_only=True):
        vals = [str(c) if c is not None else "" for c in row]
        line = " | ".join(vals)
        if line.strip(" |"):
            print(line)

# 2. BASE GENERAL SEGUIMIENTO PREVISIONAL
print("\n" + "=" * 80)
print("DOCUMENTO 2: BASE_GENERAL_SEGUIMIENTO_PREVISIONAL.xlsx")
print("=" * 80)
wb2 = openpyxl.load_workbook(os.path.join(downloads, "BASE_GENERAL_SEGUIMIENTO_PREVISIONAL.xlsx"))
for sheet_name in wb2.sheetnames:
    ws = wb2[sheet_name]
    print(f"\n--- Hoja: {sheet_name} ---")
    row_count = 0
    for row in ws.iter_rows(values_only=True):
        vals = [str(c) if c is not None else "" for c in row]
        line = " | ".join(vals)
        if line.strip(" |"):
            print(line)
            row_count += 1
            if row_count > 100:
                print(f"... (truncado, hay más filas)")
                break

# 3. especificaciones_app_centro_juridico_noa.docx
print("\n" + "=" * 80)
print("DOCUMENTO 3: especificaciones_app_centro_juridico_noa.docx")
print("=" * 80)
doc = Document(os.path.join(downloads, "especificaciones_app_centro_juridico_noa.docx"))
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# 4. PASO A PASO INGRESO JUBILACION.docx
print("\n" + "=" * 80)
print("DOCUMENTO 4: PASO A PASO INGRESO JUBILACION.docx")
print("=" * 80)
doc2 = Document(os.path.join(downloads, "PASO A PASO INGRESO JUBILACION.docx"))
for para in doc2.paragraphs:
    if para.text.strip():
        print(para.text)

# Also check tables in docx files
for doc_obj, name in [(doc, "especificaciones"), (doc2, "PASO A PASO")]:
    if doc_obj.tables:
        print(f"\n--- Tablas en {name} ---")
        for i, table in enumerate(doc_obj.tables):
            print(f"Tabla {i+1}:")
            for row in table.rows:
                vals = [cell.text for cell in row.cells]
                print(" | ".join(vals))
